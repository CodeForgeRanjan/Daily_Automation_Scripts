import streamlit as st
import pandas as pd
import re
import io
from PIL import Image
from rapidfuzz import process, utils # Fast Fuzzy matching

def play_celebration_confetti():
    """Poori screen par patakhe phodne ke liye custom JavaScript inject engine"""
    st.components.v1.html(
        """
        <script src="https://cdn.jsdelivr.net/npm/canvas-confetti@1.6.0/dist/confetti.browser.min.js"></script>
        <script>
            // Blast 1: Left Side Se
            confetti({
                particleCount: 150,
                spread: 80,
                origin: { x: 0, y: 0.6 }
            });
            // Blast 2: Right Side Se
            confetti({
                particleCount: 150,
                spread: 80,
                origin: { x: 1, y: 0.6 }
            });
            // Blast 3: Center Se Thoda Oopar
            setTimeout(() => {
                confetti({
                    particleCount: 200,
                    spread: 100,
                    origin: { y: 0.4 }
                });
            }, 300);
        </script>
        """,
        height=0, # Isse background mein silently chalega, screen par koi khali box nahi dikhega
    )

# Page Setup & Styling
st.set_page_config(page_title="Multi-Utility Automation Tool", page_icon="🚗", layout="wide")

# Custom CSS for Professional UI Design
st.markdown("""
    <style>
    .main-title {
        font-size: 40px;
        font-weight: bold;
        color: #111111;
        text-align: center;
        padding: 10px;
        background: linear-gradient(90deg, #F9D423 0%, #FF4E50 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 20px;
        border-bottom: 3px solid #FF4E50;
        border-radius: 5px;
    }
    .sidebar-heading {
        font-size: 24px;
        font-weight: bold;
        color: #00ffff;
        margin-bottom: 20px;
        text-shadow: 0px 0px 10px rgba(0,255,255,0.5);
    }
    .section-header{
        font-size: 20px;
        font-weight: 600;
        color: #F9D423;
        margin-top: 15px;
        margin-bottom: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# --- GLOBAL SHORT NAMES MAPPING DICTIONARY
CITY_ALIAS_MAP = {
    "hyd": "Hyderabad",
    "blr": "Bangalore",
    "indb": "Indore",
    "del": "Delhi",
    "cal": "Kolkata",
    "mumbai": "Mumbai",
    "chn": "Chennai",
    "pune": "Pune",
    "ahmd": "Ahmedabad"
}

# CLEANING FUNCTIONS
def clean_text_proper(text):
    if pd.isna(text) or str(text).strip() == "": 
        return ""
    cleaned = re.sub(r'[-+%,]', '', str(text)).strip()
    return cleaned.title()

def clean_address(text):
    if pd.isna(text) or str(text).strip() == "": 
        return ""
    
    text = "".join(ch for ch in str(text).strip() if ch.isprintable())


    # NEW ADVANCED HYPERLINK / DRIVE LINK REMOVAL LAYER ---
    url_pattern = r'(https?:\/\/(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b[-a-zA-Z0-9()@:%_\+.~#?&//=]*)|(www\.[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b[-a-zA-Z0-9()@:%_\+.~#?&//=]*)|(drive\.google\.com\/[^\s]*)'
    text = re.sub(url_pattern, '', text, flags=re.IGNORECASE)
    
    # If text like 'Drive Link', 'Link:', 'URL' remains in front of the link, remove it too
    text = re.sub(r'\b(?:drive\s*link|link|url)[:.-]?\s*_*', '', text, flags=re.IGNORECASE)
    

    # MOBILE NUMBER REMOVAL LAYER 
    text = re.sub(r'\b\d{5}[-\s]?\d{5}\b', '', text)
    
    unwanted_mobile_patterns = [
        r'mobile\s*no(?:umber)?[:.-]?\s*\d*',
        r'mob(?:ile)?[:.-]?\s*\d*',
        r'ph(?:one)?\s*no(?:umber)?[:.-]?\s*\d*'
    ]
    for pattern in unwanted_mobile_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    
    unwanted = [
        "Aadhar Address", "address", "Rent agreement Address", "agreement Address", 
        "AGREEMENT", "Aadhar", "AADHAR", "Aadhaar", "AADHAAR",
        "DL-", "CARD", "card", "Adhar-", "Adhar No.", "DL", "Driving License", 
        "Driving Licence", "Driving Lic", "ADD","Driving Lc", "Licence", "License", 
        "Address", "Permanent Address", "Present Address", 
        "CORRESPONDENCE ADDRESS", "CORRESPONDENCE", "PERMANENT:", "Auto_Flow" ,":", "-", ";", "#"
    ]
    for word in unwanted:
        text = re.sub(word, '', text, flags=re.IGNORECASE)
        
    return text.title().strip()

def extract_pin(text):
    if pd.isna(text) or str(text).strip() == "": 
        return ""
    
    # Get a list of all 6-digit numbers from the given address
    all_matches = re.findall(r'\d{6}', str(text))
    # If you get matches, always pick the LAST one (index -1)
    return all_matches[-1] if all_matches else ""

def split_name(name):
    parts = str(name).split()
    if len(parts) == 1: return parts[0], "", ""
    if len(parts) == 2: return parts[0], "", parts[1]
    return parts[0], " ".join(parts[1:-1]), parts[-1]

def format_id(val):
    try:
        if pd.isna(val) or str(val).strip() == "" or str(val).lower() == 'nan':
            return "NA"
        return str(int(float(val)))
    except: 
        return "NA"

def format_dob(val):
    try:
        if pd.isna(val) or str(val).strip() == "": 
            return ""
        return pd.to_datetime(val).strftime('%Y-%m-%d')
    except: 
        return str(val).replace('/', '-') 

def remove_illegal_chars(val):
    if pd.isna(val): 
        return ""
    return "".join(ch for ch in str(val) if ch.isprintable())

#  FUZZY MATCHING HELPER ENGINE 
def find_closest_city_id(detected_district, master_unique_df):
    """It will take the ID of the closest match by comparing the string from the list of all unique districts."""
    if not detected_district or detected_district == "NA":
        return "NA"
    
    master_districts = master_unique_df['DISTRICT'].dropna().astype(str).tolist()
    
    # RapidFuzz engine 
    match_result = process.extractOne(
        detected_district, 
        master_districts, 
        processor=utils.default_process,
        score_cutoff=70.0 # Will accept only if there is more than 70% similarity
    )
    if match_result:
        matched_name = match_result[0]
        # Filter the corresponding row of that matched district and get the ID
        matched_row = master_unique_df[master_unique_df['DISTRICT'].astype(str) == matched_name]
        if not matched_row.empty:
            return format_id(matched_row.iloc[0]['City ID/District ID'])
            
    return "NA"

# SIDEBAR NAVIGATION 
st.sidebar.markdown('<p class="sidebar-heading">Navigation Menu</p>', unsafe_allow_html=True)

# Aapka purana radio button menu
page = st.sidebar.radio(
    "Go to:", 
    ["Data Upload", "Image And Docs Converted", "MSG Conversion", "ARS Check Updation", "Bridge Allocation", "About Tool"]
)

# NEW PREMIUM VERSION FOOTER IN SIDEBAR ---
st.sidebar.markdown("---") 

st.sidebar.caption("🤖 **System Status:** Connected")
st.sidebar.caption("📦 **Tool Version:** `v1.2.0 (Stable)`")
st.sidebar.caption("🛡️ **Mode:** Local-First Secure")


st.sidebar.markdown(
    '<p style="color: #888888; font-size: 18px; text-align: center; margin-top: 70px;">'
    '© 2026 CodeForgeRanjan</p>', 
    unsafe_allow_html=True
)

if page == "Data Upload":
    st.markdown('<p class="main-title">Data CleanUp Dashboard</p>', unsafe_allow_html=True)
    st.write("Upload your raw CSV file and Master file to instantly generate clean data.")

    # Empty State Guide 
    st.info("Welcome! Please upload both required files below to trigger the automated data cleaning pipeline.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<p class="section-header">Raw Data Input</p>', unsafe_allow_html=True)
        my_file = st.file_uploader("Upload CSV File", type=["csv"], key="my_file", label_visibility="collapsed")

    with col2:
        st.markdown('<p class="section-header">Master Database</p>', unsafe_allow_html=True)
        master_file = st.file_uploader("Upload Pincode Master File", type=["xlsx", "xls"], key="master_file", label_visibility="collapsed")

    if my_file is not None and master_file is not None:
        st.markdown("---")
        if st.button("Reset System & Clear Cache", use_container_width=True):
            st.cache_data.clear()
            st.rerun()
        try:
            with st.spinner("Processing your data smart Fuzzy Matching... Please wait..."):
                # Load Files
                df = pd.read_csv(my_file, encoding="latin1")
                master_df = pd.read_excel(master_file, sheet_name='Sheet1')
                master_df.columns = master_df.columns.str.strip()

                # APPLY CLEANING 
                df['Candidate Name'] = df.iloc[:, 1].apply(clean_text_proper)
                df['Father Name'] = df.iloc[:, 2].apply(clean_text_proper)
                df['Cleaned_Address'] = df.iloc[:, 4].apply(clean_address)
                df['PIN_Extracted'] = df['Cleaned_Address'].apply(extract_pin)

                # Name Split logic
                df[['First', 'Middle', 'Last']] = df['Candidate Name'].apply(lambda x: pd.Series(split_name(x)))

                # MAPPING & VLOOKUP 
                df['PIN_Extracted'] = df['PIN_Extracted'].astype(str).str.strip()
                master_df['PIN CODE'] = master_df['PIN CODE'].astype(str).str.strip()
                master_unique = master_df.drop_duplicates(subset=['PIN CODE'])

                df = df.merge(
                    master_unique[['PIN CODE', 'DISTRICT', 'City ID/District ID']], 
                    left_on='PIN_Extracted', right_on='PIN CODE', how='left'
                )

                # EXACT FINAL OUTPUT STRUCTURE 
                final = pd.DataFrame()
                final['First_Name'] = df['First']
                final['Middle_Name'] = df['Middle']
                final['Last_Name'] = df['Last']
                final['Father_Name'] = df['Father Name']
                final['Mobile_No'] = ""
                final['DOB'] = df.iloc[:, 3].apply(format_dob)
                final['Location'] = "496380"
                final['Case_Insuff'] = ""
                final['Case_Comment'] = ""
                final['Car_No'] = "NOT MENTIONED"
                final['DL_NO'] = "NOT MENTIONED"
                final['Product'] = "NOT MENTIONED"
                final['UUID'] = df.iloc[:, 0]
                final['Special_ID'] = "FT_FORM"
                final['Channel'] = "OFFLINE"
                final['Permanent_Insufficiency'] = ""
                final['Name'] = ""
                final['Type'] = ""
                final['Complete_Address'] = df['Cleaned_Address']
                final['Pin_Code'] = df['PIN_Extracted']
                final['Insuff'] = ""
                final['City'] = df['DISTRICT'].fillna('NA')
                
                if 'City ID/District ID' in df.columns:
                    final['City'] = df['City ID/District ID'].apply(format_id)
                else:
                    final['City'] = "NA"       

                final['City_Name_Raw'] = df['DISTRICT'].fillna('NA')

# ADVANCED SHORT NAMES & INTELLECTUAL DISTRICT RE-CORRECTION LAYER 
                fuzzy_counter = 0
                master_district_unique = master_df.drop_duplicates(subset=['DISTRICT']).copy()
                
                for idx, row in final.iterrows():
                    # If pincode is missing initial lookup City ID 'NA'
                    if row['City'] == 'NA' or row['City'] == '' or pd.isna(row['City']):
                        address_str = str(row['Complete_Address']).lower()
                        detected_target = None
                        
                        if not detected_target:
                            address_words = re.findall(r'\b\w+\b', address_str)
                            for short_alias, full_name in CITY_ALIAS_MAP.items():
                                if short_alias in address_words:
                                    detected_target = full_name
                                    break
                        
                        # To match District names for Master Sheet
                        if not detected_target:
                            for m_dist in master_district_unique['DISTRICT'].dropna().astype(str):
                                if re.search(r'\b' + re.escape(m_dist.lower()) + r'\b', address_str):
                                    detected_target = m_dist
                                    break
                    
                        # If the target is identified by the city name of the address, apply the mapping
                        if detected_target:
                            matched_rows = master_df[master_df['DISTRICT'].astype(str).str.lower() == detected_target.lower()]
                            
                            if not matched_rows.empty:
                                # City ID update 
                                target_city_id = format_id(matched_rows.iloc[0]['City ID/District ID'])
                                target_district_name = matched_rows.iloc[0]['DISTRICT']
                                
                                final.at[idx, 'City'] = target_city_id
                                final.at[idx, 'City_Name_Raw'] = target_district_name
                                fuzzy_counter += 1
                                
                                    # Pincode Recovery: If Pincode is blank or 'NA', then enter the code of this city from the master file 
                                current_pin = str(final.at[idx, 'Pin_Code']).strip()
                                if current_pin == '' or current_pin == 'NA' or pd.isna(final.at[idx, 'Pin_Code']):
                                    # Pick up the first valid pin code of that district from the master sheet
                                    recovered_pin = str(matched_rows.iloc[0]['PIN CODE']).strip()
                                    final.at[idx, 'Pin_Code'] = recovered_pin

                final['Priority'] = ""

                # Pin Code Fallback Logic
                missing_mask = (final['City'] == 'NA') & ((final['Pin_Code'] == '') | (final['Pin_Code'].isna()))
                fallback_count = missing_mask.sum()
                final.loc[missing_mask, 'City'] = '8440'

                # Remove illegal characters from final structure
                final = final.map(remove_illegal_chars)

            st.success("Process Completed Successfully! All Data Compiled.")

            # Live Metric Cards
            m_col1, m_col2, m_col3, m_col4 = st.columns(4)
            with m_col1:
                st.metric(label="Total Input Records", value=f"{len(final)} rows")
            with m_col2:
                mapped_pins = final['Pin_Code'].transform(lambda x: 1 if x != '' else 0).sum()
                st.metric(label="Extracted Pincodes", value=f"{mapped_pins} items")
            with m_col3:
                st.metric(label="Fuzzy / City Recovery", value=f"{fuzzy_counter} rows")
            with m_col4:
                st.metric(label="Fallback Swaps (8440)", value=f"{fallback_count} rows")

# Show live preview of processed data
            st.markdown('<p class="section-header">Live Processed Preview (Top 5 Rows)</p>', unsafe_allow_html=True)
            st.dataframe(final.head(5), use_container_width=True)

            # setup for CSV stable download
            csv_buffer = io.StringIO()
            final.to_csv(csv_buffer, index=False)
            csv_output = csv_buffer.getvalue()
            
            # Stable Download 
            st.download_button(
                label=" Download Processed CSV File",
                data=csv_output,
                file_name="Uber csv (1).csv",
                mime="text/csv"
            )

        except Exception as e:
            st.error(f"error encountered during setup: {e}")
            
    elif my_file is not None and master_file is None:
        st.info(" Please upload the Master file to process automatic City & City ID mapping.")
                    
            # --- CELEBRATION BLAST START ---
            play_celebration_confetti()
            # -------------------------------

            # Live Metric Cards iske baad chalenge...



#  IMAGE & DOCS CONVERTED (ZIP + AUTO-ROTATE ENABLED) 
elif page == "Image And Docs Converted":
    import zipfile  
    from PIL import ImageOps  
    
    st.markdown('<p class="main-title">Document & Image to PDF Converter Hub</p>', unsafe_allow_html=True)
    st.write("Convert single images with auto-rotation, merge multiple images, or bulk-transform Word files into PDFs with a single 'Download All' Zip option.")

    tab1, tab2, tab3 = st.tabs(["Single Image to PDF", "Bulk Merge to One PDF", "Word Docs to PDF"])

    #  SINGLE IMAGE TO SINGLE PDF (WITH ZIP & AUTO-ROTATE) 
    with tab1:
        st.subheader("Convert Individual Images to Separate PDFs")
        single_images = st.file_uploader("Upload Images (Individual Conversion)", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="single_key")

        if single_images:
            st.success(f"Total {len(single_images)} image(s) uploaded successfully.")
            
            # If there is only 1 image then show normal download button
            if len(single_images) == 1:
                try:
                    uploaded_img = single_images[0]
                    img_data = io.BytesIO(uploaded_img.read())
                    img = Image.open(img_data)
                    
                    # Auto-Orientation 
                    img = ImageOps.exif_transpose(img)
                    img = img.convert('RGB')
                    
                    # Resolution correction to Standard A4 Layout (Vertical / Portrait)
                    img.thumbnail((1240, 1754), Image.Resampling.LANCZOS)
                    
                    pdf_buffer = io.BytesIO()
                    img.save(pdf_buffer, format="PDF")
                    pdf_bytes = pdf_buffer.getvalue()
                    
                    st.download_button(
                        label=f"Download PDF: {uploaded_img.name}.pdf",
                        data=pdf_bytes,
                        file_name=f"{uploaded_img.name.split('.')[0]}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Error: {e}")
                    
            # Show "Download All" zip button if there are more than 1 images
            else:
                if st.button("Process All Images & Create Zip", key="process_img_zip"):
                    try:
                        with st.spinner("Fixing resolution, orientation and creating Zip..."):
                            zip_buffer = io.BytesIO()
                            
                            # Create a Zip folder in memory
                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for idx, uploaded_img in enumerate(single_images):
                                    img_data = io.BytesIO(uploaded_img.read())
                                    img = Image.open(img_data)
                                    
                                    # Auto-rotate fix for flipped dimensions
                                    img = ImageOps.exif_transpose(img)
                                    img = img.convert('RGB')
                                    img.thumbnail((1240, 1754), Image.Resampling.LANCZOS)
                                    
                                    pdf_buffer = io.BytesIO()
                                    img.save(pdf_buffer, format="PDF")
                                    
                                    # Put files inside the zip
                                    clean_name = f"{uploaded_img.name.split('.')[0]}.pdf"
                                    zip_file.writestr(clean_name, pdf_buffer.getvalue())
                                    
                            st.success("All Images Converted successfully!")
                            st.download_button(
                                label="Download All PDFs in One Click (ZIP)",
                                data=zip_buffer.getvalue(),
                                file_name="All_Images_PDFs.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Zip Creation Failed: {e}")
                        
    # MULTIPLE IMAGES TO ONE MERGE PDF 
    with tab2:
        st.subheader("Compile Multiple Images into a Single Candidate PDF Report")
        bulk_images = st.file_uploader("Upload Multiple Images (All will be merged into a single PDF)", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="bulk_key")
        
        if bulk_images:
            st.success(f"Total {len(bulk_images)} images uploaded for merging.")
            if st.button("Merge All Images into 1 PDF", key="merge_btn"):
                try:
                    with st.spinner("Compiling all images with auto-rotation check..."):
                        img_list = []
                        for uploaded_img in bulk_images:
                            img_data = io.BytesIO(uploaded_img.read())
                            img = Image.open(img_data)
                            # Rotation and resolution check for merge list
                            img = ImageOps.exif_transpose(img)
                            img = img.convert('RGB')
                            img.thumbnail((1240, 1754), Image.Resampling.LANCZOS)
                            img_list.append(img)
                        
                        if img_list:
                            pdf_buffer = io.BytesIO()
                            img_list[0].save(pdf_buffer, format="PDF", save_all=True, append_images=img_list[1:])
                            pdf_data = pdf_buffer.getvalue()
                            
                            st.success("Multi-page PDF Compiled!")
                            st.download_button(
                                label="Download Compiled PDF",
                                data=pdf_data,
                                file_name="Bulk_Merged_file.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                except Exception as e:
                    st.error(f"Merge failed: {e}")

  #  UNIVERSAL WORD DOCS TO PDF 
    with tab3:
        st.markdown('<p class="section-header">Direct Word Document (.docx) to PDF Bulk Converter</p>', unsafe_allow_html=True)
        word_files = st.file_uploader("Upload Word Documents (.docx)", type=["docx"], accept_multiple_files=True, key="word_key")
        
        if word_files:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
            from reportlab.lib import colors
            
            st.success(f"{len(word_files)} docx file(s) staged for universal format conversion.")
            
            # DYNAMIC UNIVERSAL CONVERSION ENGINE 
            def universal_docx_to_pdf_story(doc, styles):
                story = []
                align_map = {0: TA_LEFT, 1: TA_CENTER, 2: TA_RIGHT, 3: TA_JUSTIFY}
                normal_style = styles['Normal']
                
                # Function that removes inner formatting of text 
                def get_clean_html_text(para):
                    text_html = ""
                    for run in para.runs:
                        text = run.text
                        if not text:
                            continue
                        # XML characters escape logic to prevent crashes
                        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        # Universal checkbox character rendering
                        text = text.replace('☐', '<font name="Helvetica" size="11">&#9633;</font>')
                        
                        if run.bold:
                            text = f"<b>{text}</b>"
                        if run.italic:
                            text = f"<i>{text}</i>"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        text_html += text
                    return text_html

                body_elements = doc.element.body
                element_index = 0
                
                for child in body_elements:
                    # PARAGRAPH PROCESSING 
                    if child.tag.endswith('p'):
                        from docx.text.paragraph import Paragraph as DocxParagraph
                        para = DocxParagraph(child, doc)
                        raw_text = para.text.strip()
                        
                        if not raw_text:
                            story.append(Spacer(1, 6))
                            continue
                        
                        # if there are large tabs/spaces between text
                        if "\t" in para.text or "   " in para.text:
                            parts = [p.strip() for p in re.split(r'\t+|\s{3,}', para.text) if p.strip()]
                            if len(parts) >= 2:
                                key_text = parts[0]
                                value_text = " ".join(parts[1:])
                                
                                k_style = ParagraphStyle(name=f'UniK_{element_index}', parent=normal_style, fontSize=9.5, leading=13, fontName="Helvetica-Bold")
                                v_style = ParagraphStyle(name=f'UniV_{element_index}', parent=normal_style, fontSize=9.5, leading=13)
                                
                                row_data = [[Paragraph(key_text, k_style), Paragraph(value_text, v_style)]]
                                kv_table = Table(row_data, colWidths=[180, 324])
                                kv_table.setStyle(TableStyle([
                                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                                    ('TOPPADDING', (0,0), (-1,-1), 3),
                                    ('BOTTOMPADDING', (0,0), (-1,-1), 3),
                                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                                ]))
                                story.append(kv_table)
                                element_index += 1
                                continue

                        text_html = get_clean_html_text(para)
                        p_align = align_map.get(para.alignment, TA_LEFT) if para.alignment is not None else TA_LEFT
                        
                        # It will dynamically scale according to the size/style 
                        p_style = ParagraphStyle(
                            name=f'UniversalPara_{element_index}',
                            parent=normal_style,
                            alignment=p_align,
                            fontSize=9.5,
                            leading=14,
                            spaceBefore=4,
                            spaceAfter=4
                        )
                        
                        story.append(Paragraph(text_html, p_style))
                        element_index += 1
                        
                    # UNIVERSAL TABLE/GRID PROCESSING
                    elif child.tag.endswith('tbl'):
                        from docx.table import Table as DocxTable
                        docx_table = DocxTable(child, doc)
                        
                        table_data = []
                        for row in docx_table.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_html = ""
                                for p in cell.paragraphs:
                                    cell_html += get_clean_html_text(p)
                                
                                cell_style = ParagraphStyle(
                                    name=f'UniCell_{element_index}_{len(row_data)}',
                                    parent=normal_style,
                                    fontSize=9,
                                    leading=12
                                )
                                row_data.append(Paragraph(cell_html, cell_style))
                            table_data.append(row_data)
                        
                        if table_data:
                            num_cols = len(table_data[0])
                            # Auto-adjust column width based on columns count 
                            col_widths = [504 / num_cols] * num_cols
                            
                            pdf_table = Table(table_data, colWidths=col_widths, repeatRows=1)
                            pdf_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F9FAFB')), 
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('TOPPADDING', (0, 0), (-1, -1), 5),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D5DB')),
                            ]))
                            story.append(Spacer(1, 6))
                            story.append(pdf_table)
                            story.append(Spacer(1, 6))
                        element_index += 1
                        
                return story

            # Single Word Document Execution
            if len(word_files) == 1:
                doc_file = word_files[0]
                if st.button(f"Convert & Process {doc_file.name}", key="single_word_btn", use_container_width=True):
                    try:
                        doc = Document(doc_file)
                        pdf_buffer = io.BytesIO()
                        doc_template = SimpleDocTemplate(
                            pdf_buffer, 
                            pagesize=letter,
                            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
                        )
                        styles = getSampleStyleSheet()
                        story = universal_docx_to_pdf_story(doc, styles)
                        doc_template.build(story)
                        
                        st.success("Document converted with universal layout styling!")
                        st.download_button(
                            label="Download Converted PDF",
                            data=pdf_buffer.getvalue(),
                            file_name=f"{doc_file.name.split('.')[0]}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Conversion Error: {e}")
            
            # Bulk Processing Archiver
            else:
                if st.button("Bulk Convert All Docs & Archive to ZIP", key="bulk_word_zip_btn", use_container_width=True):
                    try:
                        with st.spinner("Processing documents map to ZIP..."):
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for idx, doc_file in enumerate(word_files):
                                    doc = Document(doc_file)
                                    pdf_buffer = io.BytesIO()
                                    doc_template = SimpleDocTemplate(
                                        pdf_buffer, 
                                        pagesize=letter,
                                        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
                                    )
                                    styles = getSampleStyleSheet()
                                    story = universal_docx_to_pdf_story(doc, styles)
                                    doc_template.build(story)
                                    
                                    clean_name = f"{doc_file.name.split('.')[0]}.pdf"
                                    zip_file.writestr(clean_name, pdf_buffer.getvalue())
                                    
                            st.success("Bulk documents universally packed into ZIP!")
                            st.download_button(
                                label="Download Processed PDF Package (ZIP)",
                                data=zip_buffer.getvalue(),
                                file_name="All_Word_PDFs.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Bulk engine failure: {e}")

# PAGE: MSG & EML CONVERSION TO PDF SUITE 
elif page == "MSG Conversion":
        import extract_msg
        import zipfile
        from email import message_from_bytes
        from email.message import EmailMessage
        from email.utils import formatdate
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        st.markdown('<p class="main-title">Premium MSG & EML Conversion</p>', unsafe_allow_html=True)
        st.write("Upload Outlook .msg or standard .eml files to instantly unpack text structures, map attachments, and download complete compiled PDF bundles.")

        # DUAL UPLOADER BUTTONS
        col_msg, col_eml = st.columns(2)
        
        with col_msg:
            st.markdown('<p class="section-header">1: Upload msg File</p>', unsafe_allow_html=True)
            uploaded_msg = st.file_uploader("Upload Document (.msg)", type=["msg"], key="msg_file_uploader")
            
        with col_eml:
            st.markdown('<p class="section-header">2: Upload EML File</p>', unsafe_allow_html=True)
            uploaded_eml = st.file_uploader("Upload File (.eml)", type=["eml"], key="eml_file_uploader")

        # Dynamic variable initialization
        email_body_text = ""
        email_subject = "No Subject"
        email_from = "Unknown Sender"
        email_to = "Unknown Receiver"
        email_date = formatdate(localtime=True)
        valid_attachments = []
        base_name = "Email_Package"

        # IF USER UPLOADS MSG FILE 
        if uploaded_msg is not None:
            try:
                with st.spinner("Decoding Outlook MSG binary data streams..."):
                    msg = extract_msg.Message(uploaded_msg)
                    email_from = msg.sender or "Unknown Sender"
                    email_to = msg.to or "Unknown Receiver"
                    email_subject = msg.subject or "No Subject"
                    email_body_text = msg.body or ""
                    base_name = uploaded_msg.name.split('.')[0]
                    
                    # EML backup generator for download button
                    email_msg_obj = EmailMessage()
                    email_msg_obj["From"] = email_from
                    email_msg_obj["To"] = email_to
                    email_msg_obj["Subject"] = email_subject
                    email_msg_obj["Date"] = email_date
                    email_msg_obj.set_content(email_body_text)
                    eml_bytes = email_msg_obj.as_bytes()
                
                st.toast(f"EML Conversion Successful for: {base_name}",)
                st.success("Outlook Message file compiled to EML standards successfully!")
                
                st.download_button(
                    label="Download Converted EML File",
                    data=eml_bytes,
                    file_name=f"{base_name}.eml",
                    mime="message/rfc822",
                    use_container_width=True
                )
                
                # Extract attachments from MSG
                if msg.attachments:
                    for att in msg.attachments:
                        f_name = att.longFilename or att.shortFilename
                        if f_name:
                            ext = f_name.split('.')[-1].lower() if '.' in f_name else ''
                            if ext in ['pdf', 'docx']:
                                valid_attachments.append({"name": f_name, "type": ext.upper(), "raw_data": att.data})
            except Exception as e:
                st.error(f"MSG Parsing Error: {e}")

        # IF USER UPLOADS EML FILE DIRECTLY 
        elif uploaded_eml is not None:
            try:
                with st.spinner("Parsing standard EML message vectors..."):
                    eml_content = uploaded_eml.read()
                    msg_eml = message_from_bytes(eml_content)
                    
                    email_from = msg_eml.get('From', 'Unknown Sender')
                    email_to = msg_eml.get('To', 'Unknown Receiver')
                    email_subject = msg_eml.get('Subject', 'No Subject')
                    email_date = msg_eml.get('Date', formatdate(localtime=True))
                    base_name = uploaded_eml.name.split('.')[0]
                    
                    # Extract EML body content safely
                    if msg_eml.is_multipart():
                        for part in msg_eml.walk():
                            content_type = part.get_content_type()
                            content_disp = str(part.get("Content-Disposition"))
                            
                            # Extract Text body
                            if content_type == "text/plain" and "attachment" not in content_disp:
                                try:
                                    email_body_text += part.get_payload(decode=True).decode(part.get_content_charset() or 'utf-8', errors='ignore')
                                except:
                                    pass
                            # Extract Attachments from EML parts
                            elif "attachment" in content_disp or part.get_filename():
                                f_name = part.get_filename()
                                if f_name:
                                    # Handle email naming encodes if any
                                    from email.header import decode_header
                                    decoded_name = decode_header(f_name)[0][0]
                                    if isinstance(decoded_name, bytes):
                                        f_name = decoded_name.decode('utf-8', errors='ignore')
                                        
                                    ext = f_name.split('.')[-1].lower() if '.' in f_name else ''
                                    if ext in ['pdf', 'docx']:
                                        valid_attachments.append({
                                            "name": f_name, 
                                            "type": ext.upper(), 
                                            "raw_data": part.get_payload(decode=True)
                                        })
                    else:
                        email_body_text = msg_eml.get_payload(decode=True).decode(msg_eml.get_content_charset() or 'utf-8', errors='ignore')
                
                st.success(f"✨ Native EML layout verified and parsed successfully: {base_name}")
            except Exception as e:
                st.error(f"EML Parsing Error: {e}")

        # COMMON PROCESSING DISPLAY AND GENERATOR FOR BOTH CASES
        if uploaded_msg is not None or uploaded_eml is not None:
            st.markdown("---")
            st.markdown('<p class="section-header">Attachment Inventory & Metadata Tracker</p>', unsafe_allow_html=True)

            # Display Attachment Metrics UI Grid
            if valid_attachments:
                st.info(f"Found total {len(valid_attachments)} actionable document attachment(s) inside this email context.")
                att_df = pd.DataFrame(valid_attachments)[["name", "type"]]
                att_df.columns = ["File Name", "Document Type"]
                st.dataframe(att_df, use_container_width=True)
            else:
                st.warning("No valid structural document attachments (.pdf/.docx) detected inside this email structure.")

            st.markdown("---")
            
            # COMPILE BUNDLE BUTTON
            if st.button("Compile Full Email Package (Body + All Attachments) to PDF ZIP", use_container_width=True):
                with st.spinner("Executing dynamic PDF rendering algorithms... Please wait..."):
                    try:
                        zip_buffer = io.BytesIO()
                        
                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                            # Convert Email Body Text to PDF
                            body_pdf_buffer = io.BytesIO()
                            doc_template = SimpleDocTemplate(body_pdf_buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
                            styles = getSampleStyleSheet()
                            
                            story = []
                            # Email Metadata Header Block
                            story.append(Paragraph(f"<b>From:</b> {email_from}", styles['Normal']))
                            story.append(Paragraph(f"<b>To:</b> {email_to}", styles['Normal']))
                            story.append(Paragraph(f"<b>Subject:</b> {email_subject}", styles['Normal']))
                            story.append(Paragraph(f"<b>Date:</b> {email_date}", styles['Normal']))
                            story.append(Spacer(1, 15))
                            story.append(Paragraph("<b>--- EMAIL BODY TEXT ---</b>", styles['Normal']))
                            story.append(Spacer(1, 10))
                            
                            # Clean lines mapping safely
                            for line in email_body_text.split('\n'):
                                clean_line = line.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                if clean_line:
                                    story.append(Paragraph(clean_line, styles['Normal']))
                                else:
                                    story.append(Spacer(1, 6))
                                    
                            doc_template.build(story)
                            zip_file.writestr("Email_Body_Text.pdf", body_pdf_buffer.getvalue())

                            #Process and convert all attachments to PDF structure
                            for att_node in valid_attachments:
                                file_name = att_node['name']
                                file_ext = att_node['type'].lower()
                                binary_data = att_node['raw_data']
                                
                                # Attachment is already a PDF
                                if file_ext == 'pdf':
                                    zip_file.writestr(file_name, binary_data)
                                    
                                # Attachment is a Word Document (.docx)
                                elif file_ext == 'docx':
                                    from docx import Document
                                    word_stream = io.BytesIO(binary_data)
                                    doc_obj = Document(word_stream)
                                    
                                    word_pdf_buffer = io.BytesIO()
                                    word_template = SimpleDocTemplate(word_pdf_buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
                                    word_story = []
                                    
                                    for para in doc_obj.paragraphs:
                                        text_html = ""
                                        for run in para.runs:
                                            text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                            if run.bold: text = f"<b>{text}</b>"
                                            if run.italic: text = f"<i>{text}</i>"
                                            if run.underline: text = f"<u>{text}</u>"
                                            text_html += text
                                            
                                        if text_html.strip():
                                            p_style = ParagraphStyle(name=f"AttPara_{file_name}", parent=styles['Normal'], fontSize=9.5, leading=14)
                                            word_story.append(Paragraph(text_html, p_style))
                                        else:
                                            word_story.append(Spacer(1, 6))
                                            
                                    word_template.build(word_story)
                                    clean_pdf_name = f"{file_name.split('.')[0]}.pdf"
                                    zip_file.writestr(clean_pdf_name, word_pdf_buffer.getvalue())

                        # Trigger Success Layout
                        st.success("Full Package Compiled! Email body and all document templates packed securely.")
                        st.download_button(
                            label="Download Complete Converted PDF Package (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name=f"{base_name}_Full_PDF_Package.zip",
                            mime="application/zip",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Bundle Generation Error: {e}")
elif page == "ARS Check Updation":
    st.markdown('<p class="main-title"> ARS Check Updation</p>', unsafe_allow_html=True)
    st.info("Work in progress... This route is a placeholder for the background verification portal automation.")

# BRIDGE ALLOCATION SUITE 
elif page == "Bridge Allocation":
        import openpyxl
        
        st.markdown('<p class="main-title">Grafana_DATA</p>', unsafe_allow_html=True)
        st.write("Upload your raw data file to instantly clean duplicates, filter restricted series, and generate a dual-sheet allocation tracker.")

        # File Uploader supporting both CSV and Excel format
        uploaded_alloc_file = st.file_uploader("Upload Queue Data File (.csv, .xlsx)", type=["csv", "xlsx", "xls"], key="alloc_file_uploader")

        if uploaded_alloc_file is not None:
            try:
                with st.spinner("Analyzing data streams and mapping core matrices..."):
                    # DYNAMIC AUTO-DETECTING FILE LOADING LAYER 
                    if uploaded_alloc_file.name.endswith('.csv'):
                        try:
                            
                            df_alloc = pd.read_csv(uploaded_alloc_file, encoding="latin1", sep=None, engine='python', on_bad_lines='skip')
                        except Exception:
                            uploaded_alloc_file.seek(0)
                            # Backup engine metadata headers 
                            df_alloc = pd.read_csv(uploaded_alloc_file, encoding="latin1", skiprows=4, sep=None, engine='python', on_bad_lines='skip')
                    else:
                        df_alloc = pd.read_excel(uploaded_alloc_file)
                    
                    df_alloc.columns = df_alloc.columns.str.strip()
                    
                    # Tab Fallback verification 
                    if len(df_alloc.columns) <= 1:
                        uploaded_alloc_file.seek(0)
                        df_alloc = pd.read_csv(uploaded_alloc_file, encoding="latin1", sep='\t', on_bad_lines='skip')
                        df_alloc.columns = df_alloc.columns.str.strip()

                    # SAFE COLUMN IDENTIFICATION 
                    ars_candidates = [col for col in df_alloc.columns if 'ars' in col.lower() or 'ars no' in col.lower()]
                    ageing_candidates = [col for col in df_alloc.columns if 'ageing' in col.lower() or 'hour' in col.lower()]
                    
                    if ars_candidates:
                        ars_col = ars_candidates[0]
                    else:
                        st.error("File 'ARS No'! Please check columns: " + str(df_alloc.columns.tolist()))
                        st.stop()
                        
                    if ageing_candidates:
                        ageing_col = ageing_candidates[0]
                    else:
                        st.error("File 'Ageing_Hour'! Please check columns: " + str(df_alloc.columns.tolist()))
                        st.stop()
                    
                    #REMOVE DUPLICATED ARS NUMBERS 
                    initial_count = len(df_alloc)
                    df_alloc = df_alloc.dropna(subset=[ars_col])
                    df_alloc = df_alloc.drop_duplicates(subset=[ars_col])
                    dedup_count = initial_count - len(df_alloc)
                    
                    #Exclude ARS numbers starting with '2304'
                    df_alloc[ars_col] = df_alloc[ars_col].astype(str).str.strip()
                    df_filtered = df_alloc[~df_alloc[ars_col].str.startswith('2304')].copy()
                    excluded_2304_count = len(df_alloc) - len(df_filtered)
                    
                    #SORT BY AGEING HOUR (Highest ageing hours first for SLA safety)
                    df_filtered[ageing_col] = pd.to_numeric(df_filtered[ageing_col], errors='coerce').fillna(0)
                    df_filtered = df_filtered.sort_values(by=ageing_col, ascending=False).reset_index(drop=True)
                        
                    total_available_rows = len(df_filtered)

                # Show Live Queue Analytics Cards
                st.markdown('<p class="section-header">Cleaned Queue Analytics</p>', unsafe_allow_html=True)
                stat_col1, stat_col2, stat_col3 = st.columns(3)
                with stat_col1:
                    st.metric(label="Total Cases Available for Allocation", value=f"{total_available_rows} rows")
                with stat_col2:
                    st.metric(label="Duplicate Repeats Cleaned", value=f"{dedup_count} items")
                with stat_col3:
                    st.metric(label="2304 Series Blocked", value=f"{excluded_2304_count} rows")

                st.markdown("---")
                st.markdown('<p class="section-header">👥 Team Workload Allocation(5 Slots)</p>', unsafe_allow_html=True)
                st.info("Enter the User Names and the number of cases you want to allocate to each slot.")

                # Input Slots configuration
                user_allocations = []
                for idx in range(1, 6):
                    col_name, col_count = st.columns([3, 2])
                    with col_name:
                        u_name = st.text_input(f"Slot {idx}: User Name", key=f"u_name_{idx}", placeholder=f"Enter name for user {idx}...")
                    with col_count:
                        u_count = st.number_input(f"Slot {idx}: No. of Cases", min_value=0, max_value=total_available_rows, step=1, key=f"u_count_{idx}")
                    
                    if u_name.strip() != "" and u_count > 0:
                        user_allocations.append({"name": u_name.strip(), "count": int(u_count)})

                st.markdown("---")

                # TRIGGER ALLOCATION PROCESSING PIPELINE
                if st.button("Process Workload Allocation & Compile XLSX", use_container_width=True):
                    if not user_allocations:
                        st.warning("Please fill at least one User Name and a valid case count greater than 0 to allocate data!")
                    else:
                        total_requested_allocation = sum(item['count'] for item in user_allocations)
                        
                        if total_requested_allocation > total_available_rows:
                            st.error(f"Allocation Limit Exceeded! You requested total {total_requested_allocation} cases, but only {total_available_rows} cases are available.")
                        else:
                            with st.spinner("Slicing data queues and generating dual-sheet tracker framework..."):
                                current_pointer = 0
                                allocated_chunks = []
                                tracker_rows = []
                                
                                # Process slicing logic over data rows array
                                for alloc in user_allocations:
                                    name = alloc['name']
                                    count = alloc['count']
                                    
                                    # Extract the exact slice block for this user
                                    sub_df = df_filtered.iloc[current_pointer : current_pointer + count].copy()
                                    
                                    # Create a clean DataFrame with ONLY Allocated User Name and No
                                    clean_sub_df = pd.DataFrame()
                                    clean_sub_df['Allocated User Name'] = [name] * len(sub_df)
                                    clean_sub_df['ARS No'] = sub_df[ars_col].values
                                    
                                    allocated_chunks.append(clean_sub_df)
                                    
                                    # Tracking Summary row block append
                                    tracker_rows.append({
                                        "Allocated User Name": name,
                                        "Allocated Case Count": count
                                    })
                                    
                                    # Move pointer forward
                                    current_pointer += count
                                
                                # Consolidate Main Allocation Data List
                                final_allocation_df = pd.concat(allocated_chunks, ignore_index=True)
                                
                                # Create Summary Tracking DataFrame
                                final_tracker_df = pd.DataFrame(tracker_rows)
                                
                                st.success(f"Process Completed! Distributed {total_requested_allocation} cases among {len(user_allocations)} team members.")
                                
                                # UI Summary Record Display
                                st.markdown('<p class="section-header">Live Allocation Summary Record</p>', unsafe_allow_html=True)
                                st.dataframe(final_tracker_df, use_container_width=True)
                                
                                # Show preview of assigned breakdown configurations
                                st.markdown('<p class="section-header">Data Output Preview (Top 5 Rows)</p>', unsafe_allow_html=True)
                                st.dataframe(final_allocation_df.head(5), use_container_width=True)
                                
                                # Create an in-memory excel stream with DUAL SHEETS using openpyxl
                                excel_buffer = io.BytesIO()
                                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                                    # Just User Name and ARS No
                                    final_allocation_df.to_excel(writer, index=False, sheet_name='Allocation_List')
                                    # The static verification tracker counts record
                                    final_tracker_df.to_excel(writer, index=False, sheet_name='Allocation_Tracker')
                                    
                                excel_output = excel_buffer.getvalue()
                                
                                # Download button for compiled sheet
                                st.download_button(
                                    label="Download Allocated XLSX File (Dual Sheets Loaded)",
                                    data=excel_output,
                                    file_name="I_Bridge_Allocation.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True
                                )

            except Exception as e:
                st.error(f"Allocation Engine Failed: {e}")
                
elif page == "About Tool":
    st.markdown('<p class="main-title">⚙️ About Automation Utility Tool</p>', unsafe_allow_html=True)
    
    st.markdown("""
    ### 🌐 Executive Overview
    This centralized automation hub is architected under a high-performance **Offline-First Server Model** to eliminate slow, manual Excel operations and legacy macros. By replacing outdated data processes with high-speed **Python & Vectorized Pandas Pipelines**, the tool ensures 100% processing accuracy, robust data safety, and zero network-dependency bottlenecks.

    ---

    ### 🛠️ Key Production Modules & Advanced Features

    #### 📈 1. Bridge Workload Allocator (Smart Queue Distribution)
    - **Dynamic File Ingestion:** Auto-detects, reads, and cleans raw input streams seamlessly from both messy `.csv` and structured `.xlsx` workbooks.
    - **Duplicate Clean-Up Layer:** Scans incoming queues in real time, flags repetitive references, and drops duplicates to ensure exactly **1 unique case per applicant**.
    - **Compliance Series Filtering:** Instantly runs logical masks to block restricted case sequences (e.g., automatically filtering out the **`2304` series**).
    - **SLA-First Smart Sorting:** Automatically parses operational ageing strings into numeric hour indicators and sorts the entire workforce queue in **Descending Order** (highest hours first) for SLA target protection.
    - **Multi-Slot Balancing Engine:** Splits data packets dynamically into highly customized user-defined slices based on custom allocation counts.
    - **Dual-Sheet Tracker Output:** Directly compiles and outputs automated multi-sheet `.xlsx` files using an optimized `openpyxl` engine layout.

    #### 🧹 2. BGC Data CleanUp Dashboard
    - **Exact Portal Schema Mapping:** Columns are dynamically structured, re-ordered, and aligned to match production target layouts precisely.
    - **Advanced Text Sanitization:** Automatically strips illegal non-printable characters, resolves system-breaking punctuation, and applies proper Title Case.
    - **Address Deep-Cleaning Engine:** Utilizes optimized Regular Expressions (Regex) to purge clutter and junk tracking phrases from raw applicant data strings.
    - **Smart Regex PIN Extraction:** Isolate 6-digit postal pin codes instantly, even when compressed deep inside fused or unformatted text blocks.
    - **Automated Database Mapping:** Performs a high-speed memory merge against Master Databases to auto-populate exact City names and District System IDs.
    - **Conditional Fallback Logic:** Programmatically scans for rows where both Pincode and City data are unavailable, overwriting them with the system fallback code **`8440`** to guarantee 100% system ingestion success.

    #### 📄 3. Document & Image to PDF Converter Hub
    - **Multi-File Batch Processing:** Concurrently handles parallel uploads of multiple text and image formats (`.png`, `.jpg`, `.jpeg`, `.docx`).
    - **🔄 Smart Auto-Orientation (EXIF Fix):** Automatically detects if a scanned document or ID card is flipped or rotated and shifts it back to an upright portrait profile.
    - **📐 Standardized A4 Scaling:** Resizes and realigns skewed images into uniform A4 sheet boundaries using advanced LANCZOS resampling to protect text clarity.
    - **📦 One-Click Bulk ZIP Archiving:** Assembles multiple standalone file transforms directly inside memory buffers and compresses them into a single downloadable package.
    - **☁️ Cloud-Friendly Word to PDF Engine:** Uses a native hybrid parsing layer (`python-docx` & `ReportLab`) to render documents directly on the local machine without external software installations.

    ---

    ### 🛡️ Secure Core Infrastructure Strategy
    To ensure complete alignment with data safety standards, this application executes entirely within the local machine's volatile memory (RAM) and local CPU. This **Client-Side execution design** guarantees zero data exposure, ensures strict compliance with corporate Data Loss Prevention (DLP) frameworks, and allows smooth execution without external cloud endpoints.
    """)
    
